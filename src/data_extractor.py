# src/data_extractor.py

from docx import Document
import pandas as pd

def extract_data_from_word(file_path):
    extracted_data = []
    doc = Document(file_path)
    
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if '$' in paragraph.text:
                extracted_data.append(paragraph.text.strip())
    except Exception as e:
        print(f"Error processing Word document {file_path}: {str(e)}")
        
    return extracted_data


def extract_data_from_excel(file_path):
    extracted_data =[]
    try:
        df = pd.read_excel(file_path)
        extracted_data=df.to_dict(orient='records') # Convert dataframe to list of dictionaries
    except Exception as e:
        print(f"Error processing Excel file {file_path}: {str(e)}")


    return extracted_data
