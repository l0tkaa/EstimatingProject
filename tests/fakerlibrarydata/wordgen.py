import os
from dotenv import load_dotenv
from faker import Faker
from docx import Document
import random

# Create a Faker instance
fake = Faker()

# Define the directory where the Word documents will be saved
output_dir = os.getenv(WORD_SAMPLES)

# Create the directory if it does not exist
os.makedirs(output_dir, exist_ok=True)

# Function to generate a random bid amount
def generate_bid():
    return round(random.uniform(1000, 10000), 2)  # Random bid between $1,000 and $10,000

# Generate 100 fake Word documents
for i in range(1, 101):
    # Create a new Document
    doc = Document()
    
    # Add a title and some paragraphs with fake data
    doc.add_heading(f'Bid Estimate Document {i}', level=1)
    doc.add_paragraph(f"Bidder Name: {fake.name()}")
    doc.add_paragraph(f"Project: {fake.sentence(nb_words=3)}")
    doc.add_paragraph(f"Address: {fake.address()}")
    doc.add_paragraph(f"Email: {fake.email()}")
    
    # Generate a bid estimate
    bid_amount = generate_bid()
    doc.add_paragraph(f"Bid Amount: ${bid_amount:,.2f}")
    
    # Add additional estimate details
    doc.add_paragraph("Estimate Details:")
    for _ in range(3):  # Add 3 line items to the estimate
        item_description = fake.sentence(nb_words=6)
        item_cost = generate_bid()
        doc.add_paragraph(f"- {item_description}: ${item_cost:,.2f}")

    # Save the document
    doc_path = os.path.join(output_dir, f'bid_estimate_document_{i}.docx')
    doc.save(doc_path)

print("Successfully generated 100 fake Word documents with bids and estimates.")
